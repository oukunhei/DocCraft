from __future__ import annotations

import shutil
import sys
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Dict, List

try:
    from docx import Document  # type: ignore
except Exception:  # noqa: BLE001
    Document = None

try:
    from state import ReportState
except Exception:  # noqa: BLE001
    _ROOT = Path(__file__).resolve().parents[1]
    if str(_ROOT) not in sys.path:
        sys.path.insert(0, str(_ROOT))
    from state import ReportState

from nodes.docx_style import (
    StyleSpec,
    analyze_template_styles,
    append_styled_section,
    build_default_body_style,
    insert_styled_paragraph,
    replace_placeholders_preserve_format,
    set_paragraph_highlight,
)


def _resolve_template_path(template_file: str | None) -> Path:
    project_root = Path(__file__).resolve().parents[1]

    if template_file:
        p = Path(template_file)
        if p.exists():
            return p
        rel = project_root / template_file
        if rel.exists():
            return rel

    for name in ("template.docx", "templete.docx"):
        p = project_root / name
        if p.exists():
            return p

    return project_root / (template_file or "template.docx")


def _collect_slot_defs(template_slots: Dict[str, Any]) -> List[Dict[str, Any]]:
    core = template_slots.get("core_slots") or []
    other = template_slots.get("other_technical_slots") or []
    return [s for s in core + other if s.get("slot_id")]


def _fallback_filled_slots(state: ReportState) -> Dict[str, str]:
    # 1) 优先使用节点4标准输出
    analysis_notes = state.get("analysis_notes") or {}
    filled = analysis_notes.get("filled_slots") or {}
    if filled:
        out = {str(k): str(v) for k, v in filled.items() if str(v).strip()}
        if out:
            out["__source__"] = "analysis_notes"
            return out

    # 2) 回退到Agent A草稿
    out: Dict[str, str] = {}
    agent_a = state.get("agent_a_output") or {}
    a_slots = agent_a.get("slots") or []
    for item in a_slots:
        sid = str(item.get("slot_id") or "").strip()
        text = str(item.get("draft_text") or "").strip()
        if sid and text:
            out[sid] = text
    if out:
        out["__source__"] = "agent_a_output"
        return out

    # 3) 最后回退到Agent B建议文本（仅在A缺失时兜底）
    agent_b = state.get("agent_b_output") or {}
    b_slots = agent_b.get("slots") or []
    for item in b_slots:
        sid = str(item.get("slot_id") or "").strip()
        text = str(item.get("revised_text") or "").strip()
        if sid and text:
            out[sid] = text
    if out:
        out["__source__"] = "agent_b_output"
    return out


def _normalize_text(s: str) -> str:
    return "".join(ch.lower() for ch in s if ch.isalnum() or "\u4e00" <= ch <= "\u9fff")


def _find_paragraph_indices_by_title(doc: Any, title: str) -> List[int]:
    target = _normalize_text(title)
    if not target:
        return []

    indices: List[int] = []
    for idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        norm = _normalize_text(text)
        if not norm:
            continue
        # 标题匹配：相等/包含都算命中
        if norm == target or target in norm or norm in target:
            indices.append(idx)
    return indices


def _write_by_headings(
    doc: Any,
    slot_defs: List[Dict[str, Any]],
    filled_slots: Dict[str, str],
    confidence_map: Dict[str, Any],
    source_map: Dict[str, Any],
    review_flags: set[str],
    uncertainty_map: Dict[str, Any],
    style_map: Dict[str, StyleSpec],
    default_body_style: StyleSpec,
) -> Dict[str, Any]:
    inserted = 0
    unresolved: List[str] = []

    for slot in slot_defs:
        slot_id = str(slot.get("slot_id") or "").strip()
        title = str(slot.get("title") or slot_id).strip()
        content = str(filled_slots.get(slot_id, "")).strip()
        if not slot_id or not content:
            continue

        hit_indices = _find_paragraph_indices_by_title(doc, title)
        if not hit_indices:
            unresolved.append(slot_id)
            continue

        # 在首个命中标题后插入内容，使用模板分析出的样式或默认正文样式
        idx = hit_indices[0]
        spec = style_map.get(slot_id, default_body_style)
        p_content = insert_styled_paragraph(doc, idx, content, spec, parse_markdown=True)
        inserted += 1

        conf = confidence_map.get(slot_id)
        uncertainty = uncertainty_map.get(slot_id) or {}
        u_level = str(uncertainty.get("level") or "low")

        # 对低置信度/人工复核项高亮正文
        low_conf = False
        try:
            low_conf = float(conf or 0.0) < 0.6
        except Exception:  # noqa: BLE001
            low_conf = True
        if low_conf or slot_id in review_flags or u_level in {"medium", "high"}:
            set_paragraph_highlight(p_content)

    return {"inserted_by_heading": inserted, "unresolved_slots": unresolved}


def _build_review_checklist_docx(
    checklist_path: Path,
    run_id: str,
    slot_defs: List[Dict[str, Any]],
    filled_slots: Dict[str, str],
    confidence_map: Dict[str, Any],
    source_map: Dict[str, Any],
    review_flags: set[str],
    uncertainty_map: Dict[str, Any],
    default_body_style: StyleSpec,
) -> None:
    doc = Document()
    # 标题应用字体防止等线
    heading = doc.add_heading("报告复核清单", level=1)
    for run in heading.runs:
        default_body_style.apply_to_run(run)

    p = doc.add_paragraph(f"Run ID: {run_id}")
    for run in p.runs:
        default_body_style.apply_to_run(run)

    if not review_flags:
        p = doc.add_paragraph("当前没有被标记为人工复核的槽位。")
        for run in p.runs:
            default_body_style.apply_to_run(run)
        doc.save(str(checklist_path))
        return

    slot_map = {str(s.get("slot_id")): s for s in slot_defs}
    for slot_id in review_flags:
        slot = slot_map.get(slot_id, {})
        title = str(slot.get("title") or slot_id)
        content = str(filled_slots.get(slot_id, "")).strip()
        conf = confidence_map.get(slot_id)
        refs = source_map.get(slot_id, [])
        uncertainty = uncertainty_map.get(slot_id) or {}
        u_level = str(uncertainty.get("level") or "low")
        u_points = [str(x) for x in (uncertainty.get("points") or []) if str(x).strip()]
        u_human = [str(x) for x in (uncertainty.get("human_needed") or []) if str(x).strip()]

        h = doc.add_heading(f"{title} ({slot_id})", level=2)
        for run in h.runs:
            default_body_style.apply_to_run(run)

        def _add(label: str, text: str) -> None:
            para = doc.add_paragraph(f"{label}: {text}")
            for run in para.runs:
                default_body_style.apply_to_run(run)

        _add("置信度", str(conf))
        _add("引用", ", ".join(refs) if refs else "无")
        _add("不确定等级", u_level)

        if u_points:
            para = doc.add_paragraph("具体不确定点:")
            for run in para.runs:
                default_body_style.apply_to_run(run)
            for item in u_points:
                para = doc.add_paragraph(item, style="List Bullet")
                for run in para.runs:
                    default_body_style.apply_to_run(run)
        if u_human:
            para = doc.add_paragraph("需要人工补充:")
            for run in para.runs:
                default_body_style.apply_to_run(run)
            for item in u_human:
                para = doc.add_paragraph(item, style="List Bullet")
                for run in para.runs:
                    default_body_style.apply_to_run(run)

        para = doc.add_paragraph("建议检查项:")
        for run in para.runs:
            default_body_style.apply_to_run(run)
        for item in (
            "1. 结论是否可由引用证据直接支持。",
            "2. 是否存在夸大、推断过度或遗漏反例。",
            "3. 表述是否符合参赛报告风格。",
        ):
            para = doc.add_paragraph(item)
            for run in para.runs:
                default_body_style.apply_to_run(run)

        para = doc.add_paragraph("当前草稿内容:")
        for run in para.runs:
            default_body_style.apply_to_run(run)
        p = doc.add_paragraph(content if content else "[空]")
        for run in p.runs:
            default_body_style.apply_to_run(run)
        set_paragraph_highlight(p)

    doc.save(str(checklist_path))


def build_report_docx(state: ReportState) -> Dict[str, Any]:
    """节点5：基于模板创建新 docx 并填入节点4产出的槽位内容（带样式继承）。"""

    if Document is None:
        return {
            "errors": {
                "build_report_docx": "missing dependency python-docx. Please install python-docx.",
            }
        }

    template_slots = state.get("template_slots") or {}
    analysis_notes = state.get("analysis_notes") or {}
    filled_slots = _fallback_filled_slots(state)

    if not template_slots:
        return {"errors": {"build_report_docx": "missing template_slots from node2"}}
    slot_source = str(filled_slots.pop("__source__", "none")) if filled_slots else "none"

    if not filled_slots:
        return {
            "errors": {
                "build_report_docx": (
                    "missing filled slots from node4 outputs "
                    "(analysis_notes/agent_a_output/agent_b_output are empty)"
                )
            }
        }

    template_path = _resolve_template_path(state.get("template_file"))
    if not template_path.exists():
        return {
            "errors": {
                "build_report_docx": f"template_file not found: {template_path}",
            }
        }

    run_id = state.get("run_id") or datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    out_root = Path(state.get("intermediate_dir") or "artifacts/intermediate").parent / "final"
    out_root.mkdir(parents=True, exist_ok=True)
    output_path = out_root / f"report_filled_{run_id}.docx"
    checklist_path = out_root / f"report_review_checklist_{run_id}.docx"

    # 不修改原模板：先复制再写。
    shutil.copy2(template_path, output_path)

    doc = Document(str(output_path))

    # ── 样式分析 ──
    slot_defs = _collect_slot_defs(template_slots)
    slot_ids = [str(s.get("slot_id")) for s in slot_defs if s.get("slot_id")]
    style_map = analyze_template_styles(doc, slot_ids)
    default_body_style = build_default_body_style(doc)

    # 占位符替换（保留格式）
    replaced_count = replace_placeholders_preserve_format(doc, filled_slots, style_map)

    confidence_map = analysis_notes.get("slot_confidence") or {}
    source_map = analysis_notes.get("slot_sources") or {}
    review_flags = set(analysis_notes.get("review_flags") or [])
    uncertainty_map = analysis_notes.get("uncertainty_map") or {}

    # 如果 analysis_notes 不完整，回退使用 A/B 输出补齐置信度与引用。
    if not confidence_map or not source_map or not uncertainty_map:
        for item in (state.get("agent_a_output") or {}).get("slots") or []:
            sid = str(item.get("slot_id") or "").strip()
            if not sid:
                continue
            if sid not in confidence_map:
                confidence_map[sid] = item.get("confidence")
            if sid not in source_map:
                source_map[sid] = item.get("source_refs") or []
            if sid not in uncertainty_map:
                uncertainty_map[sid] = {
                    "level": "low",
                    "points": item.get("risk_notes") or [],
                    "human_needed": [],
                }
            if float(item.get("confidence") or 0.0) < 0.6:
                review_flags.add(sid)

        for item in (state.get("agent_b_output") or {}).get("slots") or []:
            sid = str(item.get("slot_id") or "").strip()
            if not sid:
                continue
            if sid not in confidence_map:
                confidence_map[sid] = item.get("confidence")
            if sid not in source_map:
                source_map[sid] = item.get("source_refs") or []
            current = uncertainty_map.get(sid) or {"level": "low", "points": [], "human_needed": []}
            points = list(current.get("points") or []) + list(item.get("disagreements") or [])
            uncertainty_map[sid] = {
                "level": "medium" if points else current.get("level", "low"),
                "points": points,
                "human_needed": current.get("human_needed") or [],
            }
            if points:
                review_flags.add(sid)

    # 第一优先：按章节标题定位写入（带样式克隆）
    heading_result = _write_by_headings(
        doc,
        slot_defs=slot_defs,
        filled_slots=filled_slots,
        confidence_map=confidence_map,
        source_map=source_map,
        review_flags=review_flags,
        uncertainty_map=uncertainty_map,
        style_map=style_map,
        default_body_style=default_body_style,
    )
    unresolved_ids = set(heading_result.get("unresolved_slots") or [])
    unresolved_slot_defs = [
        s
        for s in slot_defs
        if str(s.get("slot_id") or "") in unresolved_ids and str(filled_slots.get(str(s.get("slot_id") or ""), "")).strip()
    ]

    appended_fallback_count = 0

    # 如果模板里没有占位符，就在文档末尾追加“AI生成草稿”章节。
    if replaced_count == 0 and heading_result["inserted_by_heading"] == 0:
        appended_fallback_count = append_styled_section(
            doc=doc,
            section_title="AI 自动填充草稿（请人工复核）",
            slot_defs=slot_defs,
            filled_slots=filled_slots,
            style_spec=default_body_style,
        )
    elif unresolved_slot_defs:
        # 关键改进：部分命中模板标题时，未命中的槽位也必须保留，避免最终报告丢段。
        appended_fallback_count = append_styled_section(
            doc=doc,
            section_title="AI 自动补充草稿（未命中模板标题）",
            slot_defs=unresolved_slot_defs,
            filled_slots=filled_slots,
            style_spec=default_body_style,
        )

    doc.save(str(output_path))

    _build_review_checklist_docx(
        checklist_path=checklist_path,
        run_id=run_id,
        slot_defs=slot_defs,
        filled_slots=filled_slots,
        confidence_map=confidence_map,
        source_map=source_map,
        review_flags=review_flags,
        uncertainty_map=uncertainty_map,
        default_body_style=default_body_style,
    )

    summary = {
        "output_docx": str(output_path),
        "review_checklist_docx": str(checklist_path),
        "template_source": str(template_path),
        "slot_source": slot_source,
        "replaced_placeholders": replaced_count,
        "inserted_by_heading": heading_result["inserted_by_heading"],
        "appended_unresolved_slots": appended_fallback_count,
        "unresolved_slots": heading_result["unresolved_slots"],
        "filled_slots_count": len(filled_slots),
        "style_analyzed_slots": len(style_map),
        "default_font": default_body_style.font_name,
    }

    return {
        "run_id": run_id,
        "final_report_docx": str(output_path),
        "final_review_checklist_docx": str(checklist_path),
        "final_report": f"report generated at {output_path}",
        "report_build_summary": summary,
    }
