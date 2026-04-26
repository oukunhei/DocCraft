"""DocCraft Word 格式规范引擎

解决 python-docx 在替换/插入文本时丢失模板字体、字号、段落样式的问题。
核心思路：
1. 分析模板，提取每个 slot 对应的段落样式（StyleSpec）
2. 替换占位符时保留原 run 的字体属性
3. 插入新段落时显式克隆模板样式（段落样式 + run 字体 + 段前段后 + 行距 + 对齐）
4. 对中文东亚字体显式设置 w:eastAsia，防止退化到等线
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore
except Exception:  # noqa: BLE001
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_LINE_SPACING = None
    qn = None
    Pt = None
    RGBColor = None
    Paragraph = None


# ──────────────────────── 字体 Fallback 链 ────────────────────────

FONT_FALLBACKS: Dict[str, List[str]] = {
    "宋体": ["宋体", "SimSun", "Source Han Serif SC", "Noto Serif CJK SC"],
    "黑体": ["黑体", "SimHei", "Source Han Sans SC", "Noto Sans CJK SC"],
    "微软雅黑": ["微软雅黑", "Microsoft YaHei", "PingFang SC", "Hiragino Sans GB"],
    "仿宋": ["仿宋", "FangSong", "STFangsong"],
    "楷体": ["楷体", "KaiTi", "STKaiti"],
    "等线": ["等线", "DengXian", "Microsoft YaHei"],
}


def _resolve_font_name(name: Optional[str]) -> Optional[str]:
    """通过 fallback 链解析字体名，确保返回一个非 None 值时一定对应系统可用字体（尽力而为）。"""
    if not name:
        return None
    candidates = FONT_FALLBACKS.get(name, [name])
    return candidates[0]  # 暂不做系统存在性探测，依赖 Word 自身 fallback


# ──────────────────────── StyleSpec ────────────────────────

@dataclass
class StyleSpec:
    """从模板中提取的段落+字体样式规范。"""

    paragraph_style_name: Optional[str] = None
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    color: Optional[str] = None  # hex, e.g., "000000"
    alignment: Optional[str] = None  # LEFT, CENTER, RIGHT, JUSTIFY
    line_spacing: Optional[float] = None  # 1.5, 2.0, etc.
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    first_line_indent: bool = False

    def resolved_font(self) -> Optional[str]:
        return _resolve_font_name(self.font_name)

    @classmethod
    def from_paragraph(cls, para: Any) -> "StyleSpec":
        """从 python-docx Paragraph 对象提取样式。"""
        spec = cls()
        spec.paragraph_style_name = para.style.name if para.style else None

        # 优先从第一个 run 提取字体（显式设置的字体优先级高于样式）
        ref_run = para.runs[0] if para.runs else None
        if ref_run and ref_run.font:
            f = ref_run.font
            spec.font_name = f.name or (para.style.font.name if para.style and para.style.font else None)
            spec.bold = f.bold
            spec.italic = f.italic
            if f.size:
                spec.font_size_pt = f.size.pt if hasattr(f.size, "pt") else None
            if f.color and f.color.rgb:
                spec.color = str(f.color.rgb)
        else:
            # 从段落样式提取
            if para.style and para.style.font:
                pf = para.style.font
                spec.font_name = pf.name
                if pf.size:
                    spec.font_size_pt = pf.size.pt if hasattr(pf.size, "pt") else None
                spec.bold = pf.bold
                spec.italic = pf.italic

        # 段落格式
        pf = para.paragraph_format
        if pf.alignment is not None:
            try:
                spec.alignment = WD_ALIGN_PARAGRAPH(pf.alignment).name  # type: ignore[arg-type]
            except Exception:  # noqa: BLE001
                pass
        if pf.line_spacing is not None and isinstance(pf.line_spacing, (int, float)):
            spec.line_spacing = float(pf.line_spacing)
        if pf.space_before is not None and hasattr(pf.space_before, "pt"):
            spec.space_before_pt = pf.space_before.pt
        if pf.space_after is not None and hasattr(pf.space_after, "pt"):
            spec.space_after_pt = pf.space_after.pt
        if pf.first_line_indent is not None and hasattr(pf.first_line_indent, "pt") and pf.first_line_indent.pt > 0:
            spec.first_line_indent = True

        return spec

    def apply_to_run(self, run: Any) -> None:
        """将本 spec 的字体属性显式设置到 run 上（含东亚字体）。"""
        if qn is None:
            return
        font_name = self.resolved_font()
        if font_name:
            run.font.name = font_name
            # 关键：东亚字体必须显式设置 w:eastAsia，否则 Word 可能显示为等线
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if self.font_size_pt is not None and Pt is not None:
            run.font.size = Pt(self.font_size_pt)
        if self.bold is not None:
            run.bold = self.bold
        if self.italic is not None:
            run.italic = self.italic
        if self.color and RGBColor is not None:
            run.font.color.rgb = RGBColor.from_string(self.color)

    def apply_to_paragraph(self, para: Any, doc_styles: Optional[Any] = None) -> None:
        """将本 spec 的段落级属性设置到 paragraph 上。"""
        if self.paragraph_style_name and doc_styles is not None:
            try:
                para.style = doc_styles[self.paragraph_style_name]
            except Exception:  # noqa: BLE001
                pass

        if self.alignment and WD_ALIGN_PARAGRAPH is not None:
            try:
                para.alignment = getattr(WD_ALIGN_PARAGRAPH, self.alignment)
            except Exception:  # noqa: BLE001
                pass

        pf = para.paragraph_format
        if self.line_spacing is not None:
            pf.line_spacing = self.line_spacing
        if self.space_before_pt is not None and Pt is not None:
            pf.space_before = Pt(self.space_before_pt)
        if self.space_after_pt is not None and Pt is not None:
            pf.space_after = Pt(self.space_after_pt)
        if self.first_line_indent and self.font_size_pt is not None and Pt is not None:
            pf.first_line_indent = Pt(self.font_size_pt * 2)


# ──────────────────────── 模板样式分析 ────────────────────────

_PLACEHOLDER_PATTERNS = [("{{", "}}"), ("<<", ">>"), ("【", "】")]


def _extract_slot_ids_from_text(text: str) -> List[str]:
    """从文本中提取所有占位符里的 slot_id。"""
    found: List[str] = []
    for left, right in _PLACEHOLDER_PATTERNS:
        pattern = re.escape(left) + r"([^" + re.escape(right) + r"]+)" + re.escape(right)
        found.extend(re.findall(pattern, text))
    return found


def analyze_template_styles(doc: Any, slot_ids: List[str]) -> Dict[str, StyleSpec]:
    """分析模板文档，为每个 slot_id 提取最匹配的 StyleSpec。

    匹配优先级：
    1. 包含该 slot 占位符的段落
    2. 段落文本与 slot title 相似的段落（作为 fallback）
    """
    style_map: Dict[str, StyleSpec] = {}

    # 1) 直接命中占位符的段落
    for para in doc.paragraphs:
        text = para.text or ""
        found_ids = _extract_slot_ids_from_text(text)
        for sid in found_ids:
            if sid in slot_ids:
                style_map[sid] = StyleSpec.from_paragraph(para)

    # 2) 若还有 slot 未命中，尝试找同层级标题后的第一个正文段落作为参考样式
    if len(style_map) < len(slot_ids):
        # 扫描 Normal/Body Text 样式段落作为默认正文参考
        default_body: Optional[StyleSpec] = None
        for para in doc.paragraphs:
            name = (para.style.name or "").lower()
            if "normal" in name or "body" in name:
                default_body = StyleSpec.from_paragraph(para)
                break
        if default_body is None and doc.paragraphs:
            default_body = StyleSpec.from_paragraph(doc.paragraphs[0])

        for sid in slot_ids:
            if sid not in style_map:
                style_map[sid] = default_body or StyleSpec()

    # 3) 所有 slot 的默认 fallback：显式阻止等线退化
    for sid, spec in style_map.items():
        if not spec.font_name:
            spec.font_name = "宋体"

    return style_map


def build_default_body_style(doc: Any) -> StyleSpec:
    """构建一个默认正文样式，优先取自模板的 Normal/Body Text 段落。"""
    for para in doc.paragraphs:
        name = (para.style.name or "").lower()
        if "normal" in name or "body" in name:
            spec = StyleSpec.from_paragraph(para)
            if not spec.font_name:
                spec.font_name = "宋体"
            return spec
    if doc.paragraphs:
        spec = StyleSpec.from_paragraph(doc.paragraphs[0])
        if not spec.font_name:
            spec.font_name = "宋体"
        return spec
    return StyleSpec(font_name="宋体")


# ──────────────────────── 占位符替换（保留格式） ────────────────────────

def replace_placeholders_preserve_format(
    doc: Any,
    filled_slots: Dict[str, str],
    style_map: Optional[Dict[str, StyleSpec]] = None,
) -> int:
    """替换 docx 中的占位符，同时保留原段落的字体和样式。

    替换策略：
    - 对段落：找到包含占位符的段落，清空后重建 runs，显式复制参考 run 的字体
    - 对表格单元格：同理处理 cell.paragraphs
    """
    replaced = 0

    def _replace_in_paragraph(para: Any, slot_id: str, value: str) -> bool:
        text = para.text or ""
        for left, right in _PLACEHOLDER_PATTERNS:
            pattern = left + slot_id + right
            if pattern in text:
                # 记录参考样式
                ref_run = para.runs[0] if para.runs else None
                ref_spec = StyleSpec.from_paragraph(para)
                # 若占位符前后有文本，需要分段处理；这里简化：整段替换
                # 更精细的做法是只替换 pattern 所在 run，但为了字体一致性，
                # 我们选择清空段落、保留样式、填入新文本
                para.clear()
                run = para.add_run(value)
                # 优先用 style_map 中该 slot 的指定样式，否则继承参考 run/段落样式
                spec = (style_map or {}).get(slot_id)
                if spec is None:
                    spec = ref_spec
                spec.apply_to_run(run)
                # 段落级属性已在 clear() 后保留（style 保留），但显式再应用一次确保
                spec.apply_to_paragraph(para)
                return True
        return False

    # 处理正文段落
    for para in doc.paragraphs:
        for slot_id, value in filled_slots.items():
            if _replace_in_paragraph(para, slot_id, value):
                replaced += 1
                break  # 一个段落通常只含一个占位符

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for slot_id, value in filled_slots.items():
                        if _replace_in_paragraph(para, slot_id, value):
                            replaced += 1
                            break

    return replaced


# ──────────────────────── 插入段落（克隆样式） ────────────────────────

def insert_styled_paragraph(
    doc: Any,
    paragraph_index: int,
    text: str,
    style_spec: StyleSpec,
    parse_markdown: bool = True,
) -> Any:
    """在指定段落后插入一个新段落，并显式克隆样式。

    参数:
        doc: Document 对象
        paragraph_index: 参考段落索引（新段落插入其后）
        text: 段落文本（可含 markdown 粗体/斜体）
        style_spec: 要克隆的样式规范
        parse_markdown: 是否解析 **bold** 和 *italic*
    """
    if Document is None or qn is None:
        raise RuntimeError("python-docx not available")

    from docx.oxml import OxmlElement  # type: ignore

    ref_para = doc.paragraphs[paragraph_index]
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    new_para = Paragraph(new_p, ref_para._parent)  # type: ignore[operator]

    # 应用段落样式
    style_spec.apply_to_paragraph(new_para, doc.styles)

    # 填充文本（支持富文本）
    if parse_markdown:
        _add_rich_text(new_para, text, style_spec)
    else:
        run = new_para.add_run(text)
        style_spec.apply_to_run(run)

    return new_para


# ──────────────────────── 追加章节（带样式） ────────────────────────

def append_styled_section(
    doc: Any,
    section_title: str,
    slot_defs: List[Dict[str, Any]],
    filled_slots: Dict[str, str],
    style_spec: StyleSpec,
    title_level: int = 1,
) -> int:
    """在文档末尾追加一个带样式的章节。"""
    doc.add_page_break()
    heading = doc.add_heading(section_title, level=title_level)
    # 标题也应用字体（防止等线）
    for run in heading.runs:
        style_spec.apply_to_run(run)

    appended = 0
    for slot in slot_defs:
        sid = str(slot.get("slot_id") or "").strip()
        title = str(slot.get("title") or sid)
        content = str(filled_slots.get(sid, "")).strip()
        if not sid or not content:
            continue

        appended += 1
        h = doc.add_heading(title, level=title_level + 1)
        for run in h.runs:
            style_spec.apply_to_run(run)

        para = doc.add_paragraph()
        style_spec.apply_to_paragraph(para, doc.styles)
        _add_rich_text(para, content, style_spec)

    return appended


# ──────────────────────── 富文本支持 ────────────────────────

def _add_rich_text(para: Any, text: str, style_spec: StyleSpec) -> None:
    """解析文本中的 markdown 粗体/斜体，转为多个 run。"""
    if not text:
        return
    # 拆分 **bold** 和 *italic*，注意不处理嵌套
    tokens = re.split(r"(\*\*[^*]+?\*\*|\*[^*]+?\*)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = para.add_run(token[2:-2])
            run.bold = True
            style_spec.apply_to_run(run)
        elif token.startswith("*") and token.endswith("*"):
            run = para.add_run(token[1:-1])
            run.italic = True
            style_spec.apply_to_run(run)
        else:
            run = para.add_run(token)
            style_spec.apply_to_run(run)


# ──────────────────────── 高亮（保持兼容） ────────────────────────

try:
    from docx.enum.text import WD_COLOR_INDEX  # type: ignore
except Exception:  # noqa: BLE001
    WD_COLOR_INDEX = None


def set_paragraph_highlight(paragraph: Any) -> None:
    """对段落设置黄色高亮（兼容旧版 build_report_docx 的用法）。"""
    if WD_COLOR_INDEX is None:
        return
    if not paragraph.runs:
        run = paragraph.add_run(paragraph.text)
        paragraph.text = ""
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return
    for run in paragraph.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
